#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches

doc = Document()
doc.add_heading('OpenClaw Twitter Otomasyonları - Top 5', 0)
doc.add_paragraph(f'Hazırlayan: Claw (OpenClaw Assistant)\nTarih: 26 Şubat 2025\nPeriyot: 3 günde bir\nDil: Türkçe')

# Konu 1
doc.add_heading('1. OpenClaw + Twitter: Günlük Trend Takibi', level=1)
doc.add_paragraph('Açıklama:\nOpenClaw, Twitter Stream API\'si ile mention, hashtag ve trendleri takip eder, otomatik yanıt/ bildirim oluşturur.')
doc.add_paragraph('Nasıl Yapılır:\n- Twitter API v2 (filtered stream) kullan\n- Keywords: "OpenClaw", "Claw bot", "AI assistant"\n- Gelen tweet\'leri filter\'la → ontology\'de `Message` entity\'si oluştur\n- `twitter-reply` skill\'i ile yanıt\n- Telegram bildirimi (`message` skill)')
doc.add_paragraph('Skill\'ler:\ntwitter-stream (custom), ontology, message')
doc.add_paragraph('Fayda:\nMarka takibi, müşteri hizmetleri')

# Konu 2
doc.add_heading('2. Vibe Coding: AI ile Akış İçinde Kodlama', level=1)
doc.add_paragraph('Açıklama:\nAI asistanı ile anlık kod yazma, test, debug. Git değişikliklerini takip edip auto-review oluşturur.')
doc.add_paragraph('Nasıl Yapılır:\n- GitHub webhook → OpenClaw bildirimi\n- `github` skill\'i ile PR/commit detaylarını al\n- NVIDIA Nemotron veya Qwen modeli ile kod incelemesi\n- `ontology`\'de `Task` ve `Project` ilişkilendir\n- `nano-pdf` ile review PDF çıktısı')
doc.add_paragraph('Skill\'ler:\ngithub, openrouter (multi-model), ontology, nano-pdf')
doc.add_paragraph('Fayda:\nOtomatik kod review, zaman kazanım')

# Konu 3
doc.add_heading('3. Google Antigravity: API Kullanım Optimizasyonu', level=1)
doc.add_paragraph('Açıklama:\nBirden fazla Google API\'yi (Drive, Sheets, Gmail) tek bir optimized workflow haline getirme.')
doc.add_paragraph('Nasıl Yapılır:\n- `gog` skill\'i ile API kullanımını logla\n- Gereksiz çağrıları tespit et (rate limit)\n- Batch endpoint\'lerini kullan\n- `ontology`\'de `Account` ve `Credential` takibi')
doc.add_paragraph('Skill\'ler:\ngog, ontology, tavily')
doc.add_paragraph('Fayda:\nAPI quota tasarrufu, hız, maliyet düşüşü')

# Konu 4
doc.add_heading('4. OpenRouter Multi-Model Analiz', level=1)
doc.add_paragraph('Açıklama:\nOpenRouter\'da çoklu LLM modelleri (Qwen, Nemotron, Llama) kullanarak en uygun modeli seçmek.')
doc.add_paragraph('Nasıl Yapılır:\n- Input → text classification (model seçici)\n- `openrouter` skill\'i ile seçilen modeli çağır\n- Ensemble: sonuçları karşılaştır → en iyisi\n- `ontology`\'de `Task` ve `Document` sakla')
doc.add_paragraph('Skill\'ler:\nopenrouter, ontology, nano-pdf')
doc.add_paragraph('Fayda:\nKalite artışı, cost optimization, Türkçe performansı')

# Konu 5
doc.add_heading('5. Webhook Tabanlı Otomasyonlar', level=1)
doc.add_paragraph('Açıklama:\nFarklı servislerden (GitHub, Stripe, RSS) webhook\'ları dinleyip, OpenClaw iş akışları tetiklemek.')
doc.add_paragraph('Nasıl Yapılır:\n- HTTP server (`express` skill veya custom)\n- Webhook signature doğrulama\n- `ontology`\'de `Event` entity\'si oluştur\n- Condition → action (örn. PR açıldığında Telegram bildirimi)')
doc.add_paragraph('Skill\'ler:\nexpress (web server), ontology, message')
doc.add_paragraph('Fayda:\nEvent-driven otomasyon, entegrasyon çeşitliliği')

# Alt bilgi
doc.add_paragraph('\n---\nSonraki Adımlar:\nBu 5 otomasyon implemente edilecek. Her 3 günde bir güncelleme raporu.')

out_path = '/home/openclaw/.openclaw/workspace/raporlar/openclaw_twitter_otomasyonlari_test.docx'
doc.save(out_path)
print(f'✅ DOCX created: {out_path}')