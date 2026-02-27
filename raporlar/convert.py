#!/usr/bin/env python3
import markdown
from markdown_pdf import Extension, setup
from docx import Document
from docx.shared import Pt, Inches
import sys

md_path = sys.argv[1] if len(sys.argv)>1 else "/home/openclaw/.openclaw/workspace/raporlar/openclaw_twitter_otomasyonlari_test.md"
base = md_path.replace(".md","")

# Read markdown
with open(md_path) as f:
    html = markdown.markdown(f.read(), extensions=['extra', 'toc'])

# PDF
pdf_path = base + ".pdf"
setup(html, pdf_path, options={"displayDocTitle": True, "margin": {"top": "1cm", "bottom": "1cm", "left": "1.5cm", "right": "1.5cm"}})
print(f"✅ PDF created: {pdf_path}")

# DOCX
doc = Document()
doc.add_heading('OpenClaw Twitter Otomasyonları - Top 5', 0)
with open(md_path) as f:
    for line in f:
        line = line.rstrip()
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip() == '':
            continue
        else:
            doc.add_paragraph(line)
docx_path = base + ".docx"
doc.save(docx_path)
print(f"✅ DOCX created: {docx_path}")
