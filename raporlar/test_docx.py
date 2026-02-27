#!/usr/bin/env python3
from docx import Document
import traceback
import sys

try:
    print("🚀 DOCX oluşturuluyor...")
    doc = Document()
    doc.add_heading('OpenClaw Twitter Otomasyonları - Top 5', 0)
    doc.add_paragraph('Hazırlayan: Claw (OpenClaw Assistant)\nTarih: 26 Şubat 2025\nPeriyot: 3 günde bir\nDil: Türkçe')

    # Add just one section for test
    doc.add_heading('1. OpenClaw + Twitter: Günlük Trend Takibi', level=1)
    doc.add_paragraph('Açıklama: OpenClaw, Twitter Stream API\'si ile mention, hashtag ve trendleri takip eder.')
    doc.add_paragraph('Skill\'ler: twitter-stream, ontology, message')
    doc.add_paragraph('')

    out_path = "/tmp/test_openclaw.docx"
    doc.save(out_path)
    print(f"✅ Başarıyla oluşturuldu: {out_path}")
    print(f"   Boyut: {open(out_path,'rb').read().__len__()} bytes")
except Exception as e:
    print("❌ Hata oluştu:")
    traceback.print_exc()
    sys.exit(1)
