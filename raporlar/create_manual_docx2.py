#!/usr/bin/env python3
from docx import Document

print("🚀 DOCX oluşturuluyor...")

doc = Document()
doc.add_heading('OpenClaw Twitter Otomasyonları - Top 5', 0)
doc.add_paragraph('Hazırlayan: Claw (OpenClaw Assistant)\nTarih: 26 Şubat 2025\nPeriyot: 3 günde bir\nDil: Türkçe')

sections = [
    ("1. OpenClaw + Twitter: Günlük Trend Takibi",
     "Açıklama:\nOpenClaw, Twitter Stream API'si ile mention, hashtag ve trendleri takip eder, otomatik yanıt/ bildirim oluşturur.\n\nNasıl Yapılır:\n- Twitter API v2 (filtered stream) kullan\n- Keywords: \"OpenClaw\", \"Claw bot\", \"AI assistant\"\n- Gelen tweet'leri filter'la → ontology'de `Message` entity'si oluştur\n- `twitter-reply` skill'i ile yanıt\n- Telegram bildirimi (`message` skill)\n\nSkill'ler:\ntwitter-stream (custom), ontology, message\n\nFayda:\nMarka takibi, müşteri hizmetleri",
     ["twitter-stream (custom)", "ontology", "message"]),
    ("2. Vibe Coding: AI ile Akış İçinde Kodlama",
     "Açıklama:\nAI asistanı ile anlık kod yazma, test, debug. Git değişikliklerini takip edip auto-review oluşturur.\n\nNasıl Yapılır:\n- GitHub webhook → OpenClaw bildirimi\n- `github` skill'i ile PR/commit detaylarını al\n- NVIDIA Nemotron veya Qwen modeli ile kod incelemesi\n- `ontology`'de `Task` ve `Project` ilişkilendir\n- `nano-pdf` ile review PDF çıktısı\n\nSkill'ler:\ngithub, openrouter (multi-model), ontology, nano-pdf\n\nFayda:\nOtomatik kod review, zaman kazanım",
     ["github", "openrouter (multi-model)", "ontology", "nano-pdf"]),
    ("3. Google Antigravity: API Kullanım Optimizasyonu",
     "Açıklama:\nBirden fazla Google API'yi (Drive, Sheets, Gmail) tek bir optimized workflow haline getirme.\n\nNasıl Yapılır:\n- `gog` skill'i ile API kullanımını logla\n- Gereksiz çağrıları tespit et (rate limit)\n- Batch endpoint'lerini kullan\n- `ontology`'de `Account` ve `Credential` takibi\n\nSkill'ler:\ngog, ontology, tavily\n\nFayda:\nAPI quota tasarrufu, hız, maliyet düşüşü",
     ["gog", "ontology", "tavily"]),
    ("4. OpenRouter Multi-Model Analiz",
     "Açıklama:\nOpenRouter'da çoklu LLM modelleri (Qwen, Nemotron, Llama) kullanarak en uygun modeli seçmek.\n\nNasıl Yapılır:\n- Input → text classification (model seçici)\n- `openrouter` skill'i ile seçilen modeli çağır\n- Ensemble: sonuçları karşılaştır → en iyisi\n- `ontology`'de `Task` ve `Document` sakla\n\nSkill'ler:\nopenrouter, ontology, nano-pdf\n\nFayda:\nKalite artışı, cost optimization, Türkçe performansı",
     ["openrouter", "ontology", "nano-pdf"]),
    ("5. Webhook Tabanlı Otomasyonlar",
     "Açıklama:\nFarklı servislerden (GitHub, Stripe, RSS) webhook'ları dinleyip, OpenClaw iş akışları tetiklemek.\n\nNasıl Yapılır:\n- HTTP server (`express` skill veya custom)\n- Webhook signature doğrulama\n- `ontology`'de `Event` entity'si oluştur\n- Condition → action (örn. PR açıldığında Telegram bildirimi)\n\nSkill'ler:\nexpress (web server), ontology, message\n\nFayda:\nEvent-driven otomasyon, entegrasyon çeşitliliği",
     ["express", "ontology", "message"])
]

for title, desc, skills in sections:
    doc.add_heading(title, level=1)
    doc.add_paragraph(desc)
    doc.add_paragraph("Skill'ler:\n" + ", ".join(skills))
    doc.add_paragraph("")  # boşluk

doc.add_paragraph("\n---\nSonraki Adımlar:\nBu 5 otomasyon implemente edilecek. Her 3 günde bir güncelleme raporu.")

out_path = "/home/openclaw/.openclaw/workspace/raporlar/openclaw_twitter_otomasyonlari_test.docx"
doc.save(out_path)
print(f"✅ DOCX created: {out_path}")
